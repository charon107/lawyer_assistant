"""Tests for Stage 0: Document Preprocessor."""

import pytest
from io import BytesIO
from pathlib import Path

from app.agents.lpa.document_preprocessor import DocumentPreprocessor


class TestParseText:
    def test_parse_txt_file(self):
        preprocessor = DocumentPreprocessor()

        class FakeFile:
            name = "test.txt"
            def read(self):
                return "第一章 总则\n这是测试内容。\n第二章 出资\n出资相关内容。".encode("utf-8")

        result = preprocessor.parse(FakeFile())
        assert result["format"] == "text"
        assert result["parser_used"] == "native"
        assert "总则" in result["markdown"]
        assert result["tables"] == []
        assert result["page_count"] >= 1

    def test_parse_md_file(self):
        preprocessor = DocumentPreprocessor()

        class FakeFile:
            name = "test.md"
            def read(self):
                return "# 标题\n\n内容。".encode("utf-8")

        result = preprocessor.parse(FakeFile())
        assert result["format"] == "text"
        assert "标题" in result["markdown"]


class TestParseDocx:
    def test_parse_docx_creates_output(self):
        """Test DOCX parsing with a minimal valid docx."""
        from docx import Document

        doc = Document()
        doc.add_heading("第一章 总则", level=1)
        doc.add_paragraph("这是总则内容。")
        doc.add_heading("第二章 出资", level=1)
        doc.add_paragraph("出资相关内容。")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        preprocessor = DocumentPreprocessor()

        class FakeFile:
            name = "test.docx"
            def read(self):
                buf.seek(0)
                return buf.read()

        result = preprocessor.parse(FakeFile())
        assert result["format"] == "docx"
        assert result["parser_used"] == "python-docx"
        assert "总则" in result["raw_text"]


class TestBuildOutput:
    def test_output_structure(self):
        preprocessor = DocumentPreprocessor()
        result = preprocessor._build_output("test text", [], "text", "native")
        assert "markdown" in result
        assert "tables" in result
        assert "page_count" in result
        assert "format" in result
        assert "parser_used" in result
        assert "raw_text" in result

    def test_tables_preserved(self):
        tables = [[["A", "B"], ["1", "2"]]]
        preprocessor = DocumentPreprocessor()
        result = preprocessor._build_output("text", tables, "digital_pdf", "pypdf")
        assert result["tables"] == tables


class TestTextToMarkdown:
    def test_removes_extra_blank_lines(self):
        text = "line1\n\n\n\nline2"
        md = DocumentPreprocessor._text_to_markdown(text)
        assert "\n\n\n" not in md

    def test_strips_whitespace(self):
        text = "  line1  \n  line2  "
        md = DocumentPreprocessor._text_to_markdown(text)
        assert md == "line1\nline2"

    def test_empty_text(self):
        md = DocumentPreprocessor._text_to_markdown("")
        assert md == ""


class TestEstimatePages:
    def test_chinese_text(self):
        text = "中" * 1000
        pages = DocumentPreprocessor._estimate_pages(text)
        assert pages == 2

    def test_english_text(self):
        text = "a" * 5000
        pages = DocumentPreprocessor._estimate_pages(text)
        assert pages == 2

    def test_minimum_one_page(self):
        assert DocumentPreprocessor._estimate_pages("") == 1
        assert DocumentPreprocessor._estimate_pages("short") == 1


class TestDetectScanned:
    def test_scanned_detection(self):
        preprocessor = DocumentPreprocessor()
        assert preprocessor._detect_scanned("") is True
        assert preprocessor._detect_scanned("   ") is True
        assert preprocessor._detect_scanned("a" * 200) is False


class TestUnsupportedFormat:
    def test_raises_on_unsupported(self):
        preprocessor = DocumentPreprocessor()

        class FakeFile:
            name = "test.xyz"
            def read(self):
                return b"data"

        with pytest.raises(ValueError, match="Unsupported file type"):
            preprocessor.parse(FakeFile())
